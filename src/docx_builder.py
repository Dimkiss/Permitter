import random

from docx import Document
from docx.shared import Pt


class DocxBuilder:
    PLACEHOLDERS = {
        "{{DATE}}": "date",
        "{{TIME_FROM}}": "time_from",
        "{{TIME_TO}}": "time_to",
        "{{LOCATION}}": "location",
        "{{RESPONSIBLE_NAME}}": "responsible_name",
        "{{PHONE}}": "phone",
        "{{OPERATION_TYPE}}": "operation_type",
    }

    def replace_in_paragraph(self, paragraph, data):
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return

        updated_text = full_text
        for placeholder, key in self.PLACEHOLDERS.items():
            updated_text = updated_text.replace(placeholder, str(data.get(key, "")))

        if updated_text != full_text:
            for run in paragraph.runs:
                run.text = ""

            if paragraph.runs:
                paragraph.runs[0].text = updated_text
            else:
                paragraph.add_run(updated_text)

    def replace_everywhere(self, doc, data):
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, data)

    def generate_inventory_number(self):
        return str(random.randint(100000, 999999))

    def set_cell_text(self, cell, text, font_size=10):
        cell.text = str(text)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)

    def write_items_table(self, doc, items):
        if not doc.tables:
            raise ValueError("В шаблоне нет таблицы для списка материальных ценностей.")

        table = doc.tables[0]

        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        for index, item in enumerate(items, start=1):
            row = table.add_row().cells

            self.set_cell_text(row[0], index)
            self.set_cell_text(row[1], item["name"])
            self.set_cell_text(row[2], item.get("quantity", ""))
            self.set_cell_text(row[3], self.generate_inventory_number())

            if len(row) > 4:
                self.set_cell_text(row[4], "")

    def build(self, template_path, output_path, data, items):
        doc = Document(template_path)
        self.replace_everywhere(doc, data)
        self.write_items_table(doc, items)
        doc.save(output_path)