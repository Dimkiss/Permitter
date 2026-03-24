import os
import re
import traceback
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import pandas as pd
from docx import Document
from docx.shared import Pt


def normalize_text(value):
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").strip()
    return re.sub(r"\s+", " ", text)


class OutlookMailer:
    def __init__(self):
        self.outlook = None

    def check(self):
        try:
            import pythoncom
            import win32com.client

            pythoncom.CoInitialize()
            self.outlook = win32com.client.Dispatch("Outlook.Application")
            session = self.outlook.Session

            if session.Accounts.Count == 0:
                return False, "Outlook найден, но ни один аккаунт не подключён."

            accounts = []
            for i in range(1, session.Accounts.Count + 1):
                acc = session.Accounts.Item(i)
                address = getattr(acc, "SmtpAddress", "") or getattr(acc, "DisplayName", "")
                accounts.append(address)

            return True, f"Outlook готов. Аккаунты: {', '.join(accounts)}"
        except Exception as e:
            return False, f"Outlook не установлен или недоступен: {e}"

    def send(self, to_email, subject, body, attachment_path, display_only=False):
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        outlook = win32com.client.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        mail.To = to_email
        mail.Subject = subject
        mail.Body = body
        if attachment_path:
            mail.Attachments.Add(str(Path(attachment_path).resolve()))
        if display_only:
            mail.Display()
        else:
            mail.Send()


class InvoiceParser:
    PRODUCT_HEADER_KEYWORDS = [
        "наименование товара",
        "наименование продукции",
        "товар",
        "наименование",
        "работ, услуг",
    ]
    QTY_HEADER_KEYWORDS = ["количество", "кол-во", "кол во", "qty"]
    UNIT_HEADER_KEYWORDS = ["ед. изм", "ед изм", "единица", "ед."]

    def read_excel(self, path):
        ext = Path(path).suffix.lower()
        if ext == ".xls":
            return pd.read_excel(path, header=None, engine="xlrd")
        if ext == ".xlsx":
            return pd.read_excel(path, header=None, engine="openpyxl")
        raise ValueError("Поддерживаются только .xls и .xlsx")

    def find_header_row(self, df):
        for idx in range(len(df)):
            row_values = [normalize_text(x).lower() for x in df.iloc[idx].tolist()]
            row_text = " | ".join(row_values)
            if any(k in row_text for k in self.PRODUCT_HEADER_KEYWORDS) and any(
                k in row_text for k in self.QTY_HEADER_KEYWORDS
            ):
                return idx
        return None

    def find_column(self, headers, keywords):
        for i, header in enumerate(headers):
            value = normalize_text(header).lower()
            if any(k in value for k in keywords):
                return i
        return None

    def is_summary_row(self, text):
        text = text.lower()
        markers = [
            "итого",
            "всего к оплате",
            "в том числе",
            "сумма",
            "ндс",
            "всего наименований",
        ]
        return any(m in text for m in markers)

    def parse_items(self, path):
        df = self.read_excel(path)
        header_row = self.find_header_row(df)
        if header_row is None:
            raise ValueError(
                "Не удалось найти строку заголовков. Нужны колонки с наименованием и количеством."
            )

        headers = df.iloc[header_row].tolist()
        name_col = self.find_column(headers, self.PRODUCT_HEADER_KEYWORDS)
        qty_col = self.find_column(headers, self.QTY_HEADER_KEYWORDS)
        unit_col = self.find_column(headers, self.UNIT_HEADER_KEYWORDS)

        if name_col is None or qty_col is None:
            raise ValueError("Не удалось определить колонки с наименованием и количеством.")

        items = []
        for idx in range(header_row + 1, len(df)):
            row = df.iloc[idx].tolist()
            if name_col >= len(row):
                continue

            name = normalize_text(row[name_col])
            if not name or name.lower() == "nan":
                continue
            if self.is_summary_row(name):
                break

            qty = row[qty_col] if qty_col < len(row) else ""
            unit = row[unit_col] if unit_col is not None and unit_col < len(row) else ""

            items.append(
                {
                    "name": name,
                    "qty": normalize_text(qty),
                    "unit": normalize_text(unit),
                    "serial": "-",
                    "note": "",
                }
            )

        if not items:
            raise ValueError("Товары не найдены. Проверь структуру счёта.")

        return items


class DocxBuilder:
    PLACEHOLDERS = {
        "{{DATE}}": "date",
        "{{TIME_FROM}}": "time_from",
        "{{TIME_TO}}": "time_to",
        "{{LOCATION_FROM}}": "location_from",
        "{{LOCATION_TO}}": "location_to",
        "{{RESPONSIBLE_NAME}}": "responsible_name",
        "{{PHONE}}": "phone",
        "{{OPERATION_TYPE}}": "operation_type",
        "{{COMMENT}}": "comment",
    }

    def replace_in_paragraph(self, paragraph, data):
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return

        updated_text = full_text
        for placeholder, key in self.PLACEHOLDERS.items():
            updated_text = updated_text.replace(placeholder, str(data.get(key, "")))

        if updated_text != full_text:
            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph.runs[i].text = ""
            if paragraph.runs:
                paragraph.runs[0].text = updated_text
            else:
                paragraph.add_run(updated_text)

    def replace_everywhere(self, doc, data):
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, data)

    def write_items_table(self, doc, items):
        if not doc.tables:
            raise ValueError("В шаблоне нет таблицы для списка материальных ценностей.")

        table = doc.tables[0]
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        for index, item in enumerate(items, start=1):
            row = table.add_row().cells
            row[0].text = str(index)
            row[1].text = item["name"]
            row[2].text = str(item.get("serial", "-")) or "-"
            row[3].text = f"{item.get('qty', '')} {item.get('unit', '')}".strip()
            row[4].text = item.get("note", "")

            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    def build(self, template_path, output_path, data, items):
        doc = Document(template_path)
        self.replace_everywhere(doc, data)
        self.write_items_table(doc, items)
        doc.save(output_path)


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("RPA: разрешение на перемещение МЦ")
        self.root.geometry("900x720")

        self.parser = InvoiceParser()
        self.builder = DocxBuilder()
        self.mailer = OutlookMailer()
        self.items = []

        self.invoice_path = tk.StringVar()
        self.template_path = tk.StringVar()
        self.output_dir = tk.StringVar(value=str(Path.cwd()))
        self.email_to = tk.StringVar(value="security@example.com")
        self.subject = tk.StringVar(value="Разрешение на перемещение материальных ценностей")
        self.operation_type = tk.StringVar(value="Внос")
        self.date_value = tk.StringVar(value=datetime.now().strftime("%d.%m.%Y"))
        self.time_from = tk.StringVar(value="09:00")
        self.time_to = tk.StringVar(value="18:00")
        self.location_from = tk.StringVar(value="Склад компании")
        self.location_to = tk.StringVar(value="Рабочее помещение заказчика")
        self.responsible_name = tk.StringVar(value="Иванов И.И.")
        self.phone = tk.StringVar(value="+7 900 000-00-00")
        self.comment = tk.StringVar(value="")
        self.preview_before_send = tk.BooleanVar(value=True)

        self.build_ui()

    def log(self, text):
        self.log_box.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')}  {text}\n")
        self.log_box.see(tk.END)
        self.root.update_idletasks()

    def build_ui(self):
        top = ttk.Frame(self.root, padding=12)
        top.pack(fill="both", expand=True)

        files_frame = ttk.LabelFrame(top, text="Файлы", padding=10)
        files_frame.pack(fill="x", pady=(0, 10))

        ttk.Label(files_frame, text="Счёт Excel (.xls/.xlsx)").grid(row=0, column=0, sticky="w")
        ttk.Entry(files_frame, textvariable=self.invoice_path, width=80).grid(row=1, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(files_frame, text="Выбрать", command=self.pick_invoice).grid(row=1, column=1)

        ttk.Label(files_frame, text="Шаблон Word (.docx)").grid(row=2, column=0, sticky="w", pady=(8, 0))
        ttk.Entry(files_frame, textvariable=self.template_path, width=80).grid(row=3, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(files_frame, text="Выбрать", command=self.pick_template).grid(row=3, column=1)

        ttk.Label(files_frame, text="Папка для результата").grid(row=4, column=0, sticky="w", pady=(8, 0))
        ttk.Entry(files_frame, textvariable=self.output_dir, width=80).grid(row=5, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(files_frame, text="Выбрать", command=self.pick_output_dir).grid(row=5, column=1)
        files_frame.columnconfigure(0, weight=1)

        form_frame = ttk.LabelFrame(top, text="Данные разрешения", padding=10)
        form_frame.pack(fill="x", pady=(0, 10))

        labels = [
            ("Тип операции", self.operation_type),
            ("Дата", self.date_value),
            ("Время с", self.time_from),
            ("Время до", self.time_to),
            ("Точка выноса", self.location_from),
            ("Точка вноса", self.location_to),
            ("Ответственное лицо", self.responsible_name),
            ("Телефон", self.phone),
            ("Комментарий", self.comment),
            ("Email получателя", self.email_to),
            ("Тема письма", self.subject),
        ]

        for i, (label, variable) in enumerate(labels):
            ttk.Label(form_frame, text=label).grid(row=i, column=0, sticky="w", pady=3)
            if variable is self.operation_type:
                combo = ttk.Combobox(form_frame, textvariable=variable, values=["Внос", "Вынос", "Перемещение"], state="readonly", width=40)
                combo.grid(row=i, column=1, sticky="ew", pady=3)
            else:
                ttk.Entry(form_frame, textvariable=variable, width=60).grid(row=i, column=1, sticky="ew", pady=3)
        form_frame.columnconfigure(1, weight=1)

        options_frame = ttk.Frame(top)
        options_frame.pack(fill="x", pady=(0, 10))
        ttk.Checkbutton(options_frame, text="Показывать письмо перед отправкой", variable=self.preview_before_send).pack(anchor="w")

        buttons_frame = ttk.Frame(top)
        buttons_frame.pack(fill="x", pady=(0, 10))
        ttk.Button(buttons_frame, text="1. Проверить Outlook", command=self.check_outlook).pack(side="left", padx=(0, 8))
        ttk.Button(buttons_frame, text="2. Прочитать Excel", command=self.parse_invoice).pack(side="left", padx=(0, 8))
        ttk.Button(buttons_frame, text="3. Создать Word", command=self.build_docx).pack(side="left", padx=(0, 8))
        ttk.Button(buttons_frame, text="4. Отправить Email", command=self.send_email).pack(side="left")

        preview_frame = ttk.LabelFrame(top, text="Найденные позиции", padding=10)
        preview_frame.pack(fill="both", expand=True, pady=(0, 10))

        columns = ("name", "qty", "unit")
        self.tree = ttk.Treeview(preview_frame, columns=columns, show="headings", height=10)
        self.tree.heading("name", text="Наименование")
        self.tree.heading("qty", text="Количество")
        self.tree.heading("unit", text="Ед.")
        self.tree.column("name", width=540)
        self.tree.column("qty", width=120, anchor="center")
        self.tree.column("unit", width=100, anchor="center")
        self.tree.pack(fill="both", expand=True)

        log_frame = ttk.LabelFrame(top, text="Лог", padding=10)
        log_frame.pack(fill="both", expand=True)
        self.log_box = tk.Text(log_frame, height=10)
        self.log_box.pack(fill="both", expand=True)

    def pick_invoice(self):
        path = filedialog.askopenfilename(filetypes=[("Excel", "*.xls *.xlsx")])
        if path:
            self.invoice_path.set(path)

    def pick_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if path:
            self.template_path.set(path)

    def pick_output_dir(self):
        path = filedialog.askdirectory()
        if path:
            self.output_dir.set(path)

    def check_outlook(self):
        ok, message = self.mailer.check()
        self.log(message)
        if ok:
            messagebox.showinfo("Outlook", message)
        else:
            messagebox.showwarning("Outlook", message)

    def parse_invoice(self):
        path = self.invoice_path.get().strip()
        if not path:
            messagebox.showwarning("Ошибка", "Выбери Excel-файл.")
            return

        try:
            self.log("Чтение Excel...")
            self.items = self.parser.parse_items(path)
            for row in self.tree.get_children():
                self.tree.delete(row)
            for item in self.items:
                self.tree.insert("", tk.END, values=(item["name"], item["qty"], item["unit"]))
            self.log(f"Найдено позиций: {len(self.items)}")
        except Exception as e:
            self.log(f"Ошибка разбора Excel: {e}")
            messagebox.showerror("Ошибка", str(e))

    def collect_doc_data(self):
        return {
            "operation_type": self.operation_type.get().strip(),
            "date": self.date_value.get().strip(),
            "time_from": self.time_from.get().strip(),
            "time_to": self.time_to.get().strip(),
            "location_from": self.location_from.get().strip(),
            "location_to": self.location_to.get().strip(),
            "responsible_name": self.responsible_name.get().strip(),
            "phone": self.phone.get().strip(),
            "comment": self.comment.get().strip(),
        }

    def get_output_docx_path(self):
        safe_date = datetime.now().strftime("%Y%m%d_%H%M%S")
        return str(Path(self.output_dir.get()) / f"permit_{safe_date}.docx")

    def build_docx(self):
        template_path = self.template_path.get().strip()
        if not template_path:
            messagebox.showwarning("Ошибка", "Выбери шаблон Word.")
            return
        if not self.items:
            messagebox.showwarning("Ошибка", "Сначала прочитай Excel.")
            return

        try:
            output_path = self.get_output_docx_path()
            self.log("Формирование Word...")
            self.builder.build(template_path, output_path, self.collect_doc_data(), self.items)
            self.log(f"Документ сохранён: {output_path}")
            messagebox.showinfo("Готово", f"Документ создан:\n{output_path}")
        except Exception as e:
            self.log(f"Ошибка формирования Word: {e}")
            self.log(traceback.format_exc())
            messagebox.showerror("Ошибка", str(e))

    def send_email(self):
        if not self.items:
            messagebox.showwarning("Ошибка", "Сначала прочитай Excel.")
            return
        template_path = self.template_path.get().strip()
        if not template_path:
            messagebox.showwarning("Ошибка", "Выбери шаблон Word.")
            return

        try:
            output_path = self.get_output_docx_path()
            if not Path(output_path).exists():
                self.log("Создание Word перед отправкой...")
                self.builder.build(template_path, output_path, self.collect_doc_data(), self.items)

            self.log("Проверка Outlook...")
            ok, msg = self.mailer.check()
            self.log(msg)
            if not ok:
                messagebox.showwarning("Outlook", msg)
                return

            self.log("Создание письма...")
            body = (
                f"Добрый день.\n\n"
                f"Направляю разрешение на {self.operation_type.get().lower()} материальных ценностей.\n"
                f"Дата: {self.date_value.get()}\n"
                f"Время: {self.time_from.get()} - {self.time_to.get()}\n"
                f"Ответственное лицо: {self.responsible_name.get()}\n"
                f"Телефон: {self.phone.get()}\n\n"
                f"Документ во вложении."
            )
            self.mailer.send(
                to_email=self.email_to.get().strip(),
                subject=self.subject.get().strip(),
                body=body,
                attachment_path=output_path,
                display_only=self.preview_before_send.get(),
            )
            self.log("Письмо успешно сформировано/отправлено.")
            messagebox.showinfo("Готово", "Письмо успешно сформировано/отправлено.")
        except Exception as e:
            self.log(f"Ошибка отправки: {e}")
            self.log(traceback.format_exc())
            messagebox.showerror("Ошибка", str(e))


def main():
    root = tk.Tk()
    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
